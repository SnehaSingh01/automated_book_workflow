from scraping.scraper import scrape_chapter
from ai_agents.writer import spin_text
from ai_agents.reviewer import review_text
from ai_agents.editor_interface import human_review
from versioning.chroma_db import store_version, retrieve_version, load_all_versions
from rl_search.rl_search import rl_search
from docx import Document
from fpdf import FPDF
import os


def export_to_docx(text, filename):
    doc = Document()
    doc.add_heading(' Final Edited Chapter', 0)
    for paragraph in text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(filename)
    print(f" DOCX saved at: {filename}")


def export_to_pdf(text, filename):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for line in text.split('\n'):
        pdf.multi_cell(0, 10, txt=line)
    pdf.output(filename)
    print(f" PDF saved at: {filename}")


def run_pipeline():
    print("\n Running Automated Book Pipeline...\n")

    # Step 1: Scrape Chapter
    url = "https://en.wikisource.org/wiki/The_Gates_of_Morning/Book_1/Chapter_1"
    content = scrape_chapter(url)
    print(" Scraped content.")

    # Step 2: AI Writer
    spun = spin_text(content)
    print(" AI-generated draft complete.")

    # Step 3: AI Reviewer
    reviewed = review_text(spun)
    print(" AI-reviewed draft complete.")

    # Step 4: Human-in-the-loop Editing
    final = human_review(reviewed)
    print(" Human-reviewed final version ready.")

    # Ensure output folder exists
    os.makedirs("output", exist_ok=True)

    # Step 5: Save reviewed text for UI
    reviewed_path = "output/ai_reviewed_chapter.txt"
    with open(reviewed_path, "w", encoding="utf-8") as f:
        f.write(final)
    print(f" Saved AI-reviewed text at: {reviewed_path}")

    # Step 6: Store Final Version in ChromaDB
    version_id = "chapter1_final"
    store_version(version_id, final)
    print(" Stored in ChromaDB with version ID:", version_id)

    # Step 7: Export to DOCX and PDF
    export_to_docx(final, "output/final_chapter.docx")
    export_to_pdf(final, "output/final_chapter.pdf")

    # Step 8: RL Search Simulation
    query = input("\n Enter a search query to find relevant version (e.g., 'war and defense'): ")
    all_versions = load_all_versions()
    best_match_id = rl_search(query, all_versions)
    best_content = retrieve_version(best_match_id)

    print(f"\n Best match found: {best_match_id}")
    print(" Snippet of matched content:\n")
    print(best_content[:800], "...\n")


if __name__ == "__main__":
    run_pipeline()
